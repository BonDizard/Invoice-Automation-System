import os
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import tkinter.font as tkfont
import ctypes
import docx
from docx2pdf import convert
from PIL import Image, ImageTk
import subprocess
import datetime as dt
import tempfile

class InvoiceAutomation:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title('Invoice Automation')
        self.root.geometry('500x700')
        self.root.configure(bg='#f0f2f5')

        # Load custom font
        font_path = os.path.abspath("assets/PokemonGb-RAeo.ttf")
        self.register_font(font_path)

        try:
            self.custom_font = tkfont.Font(family="Pokemon GB", size=11)
        except tk.TclError:
            self.custom_font = tkfont.Font(family="Segoe UI", size=11)

        # Main container
        self.main_frame = tk.Frame(self.root, bg='white', bd=2, relief='flat')
        self.main_frame.pack(padx=20, pady=20, fill='both', expand=True)

        # Title
        self.title_label = tk.Label(
            self.main_frame,
            text='Invoice Automation',
            font=("Segoe UI", 16, "bold"),
            bg='white',
            fg='#333333'
        )
        self.title_label.pack(pady=15)

        # Form frame
        self.form_frame = tk.Frame(self.main_frame, bg='white')
        self.form_frame.pack(padx=20, pady=10, fill='x')

        # Form fields
        self.create_form_field("Date", "date")
        self.create_form_field("Reference Number", "reference_number")
        self.create_form_field("Name", "name")
        self.create_form_field("Amount", "amount")
        self.create_form_field("Amount in Words", "amount_words")

        # Set default value for date
        self.date_entry.insert(0, dt.datetime.today().strftime('%d/%m/%Y'))
        
        # Auto-update amount in words when amount changes
        self.amount_entry.bind('<KeyRelease>', self.update_amount_words)

        # Payment mode dropdown
        self.payment_mode = tk.StringVar(self.form_frame)
        self.payment_mode.set('NEFT')

        style = ttk.Style()
        style.configure(
            "Custom.TMenubutton",
            font=self.custom_font,
            background='#ffffff',
            foreground='#333333',
            borderwidth=1,
            relief='solid',
            padding=6
        )
        style.map(
            "Custom.TMenubutton",
            background=[('active', '#e6e6e6')],
            relief=[('pressed', 'sunken')]
        )

        self.payment_mode_dropdown = ttk.OptionMenu(
            self.form_frame,
            self.payment_mode,
            'NEFT',
            'NEFT',
            'RTGS',
            'CHEQUE',
            style="Custom.TMenubutton"
        )
        self.payment_mode_dropdown.pack(fill='x', padx=5, pady=10)
        menu = self.payment_mode_dropdown['menu']
        menu.config(font=self.custom_font, bg='white', fg='#333333')

        # Create Invoice Button
        self.create_button = tk.Button(
            self.main_frame,
            text='Create Invoice',
            font=("Segoe UI", 12, "bold"),
            bg='#4CAF50',
            fg='white',
            bd=0,
            relief='flat',
            command=self.create_invoice,
            cursor='hand2',
            activebackground='#45a049',
            activeforeground='white',
            padx=20,
            pady=10
        )
        self.create_button.pack(pady=20)
        self.create_button.bind("<Enter>", lambda e: self.create_button.config(bg='#45a049'))
        self.create_button.bind("<Leave>", lambda e: self.create_button.config(bg='#4CAF50'))

        # Static image
        self.image_path = r"assets\poke_ball.png"  # Replace with your static image path
        try:
            if os.path.exists(self.image_path):
                image = Image.open(self.image_path)
                image = image.resize((50, 50), Image.LANCZOS)
                self.static_image = ImageTk.PhotoImage(image)
                self.placeholder_image = tk.Label(self.main_frame, image=self.static_image, bg='white')
                self.placeholder_image.pack(pady=10)
            else:
                raise FileNotFoundError("Static image file not found")
        except Exception as e:
            print(f"Image Error: {e}")
            self.placeholder_image = tk.Label(
                self.main_frame,
                text="[Failed to load pikachu_pokemon_2_28_24.jpg]",
                bg='white',
                fg='#888888',
                font=("Segoe UI", 10)
            )
            self.placeholder_image.pack(pady=10)

        self.root.mainloop()

    def create_form_field(self, label_text, field_name):
        """Helper method to create a labeled entry field."""
        label = tk.Label(
            self.form_frame,
            text=label_text,
            font=self.custom_font,
            bg='white',
            fg='#333333'
        )
        label.pack(anchor='w', padx=5, pady=2)

        entry = tk.Entry(
            self.form_frame,
            font=self.custom_font,
            bd=1,
            relief='solid',
            bg='#f9f9f9',
            fg='#333333',
            highlightthickness=1,
            highlightbackground='#dcdcdc',
            highlightcolor='#4CAF50'
        )
        entry.pack(fill='x', padx=5, pady=5)
        setattr(self, f"{field_name}_entry", entry)

    def register_font(self, font_path):
        """Registers a .ttf font on Windows at runtime using ctypes."""
        if os.name == 'nt' and os.path.exists(font_path):
            ctypes.windll.gdi32.AddFontResourceW(font_path)
            ctypes.windll.user32.SendMessageW(0xFFFF, 0x001D, 0, font_path)
        else:
            messagebox.showerror("Font Error", f"Font file not found: {font_path}")

    def update_amount_words(self, event=None):
        """Update amount in words field when amount changes."""
        try:
            amount = float(self.amount_entry.get())
            words = self.convert_to_words(amount)
            self.amount_words_entry.delete(0, tk.END)
            self.amount_words_entry.insert(0, f"Indian Rupee {words}")
        except ValueError:
            self.amount_words_entry.delete(0, tk.END)
            self.amount_words_entry.insert(0, "Indian Rupee")

    @staticmethod
    def convert_to_words(amount):
        """Convert numeric amount to words in Indian Rupee format."""
        if not amount:
            return "Zero Only"

        units = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine"]
        teens = ["Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen", 
                 "Seventeen", "Eighteen", "Nineteen"]
        tens = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]
        thousands = ["", "Thousand", "Lakh", "Crore"]

        def convert_below_thousand(num):
            if num == 0:
                return ""
            elif num < 10:
                return units[num]
            elif num < 20:
                return teens[num - 10]
            elif num < 100:
                return f"{tens[num // 10]} {units[num % 10]}".strip()
            else:
                return f"{units[num // 100]} Hundred {'and ' + convert_below_thousand(num % 100) if num % 100 else ''}".strip()

        amount = round(float(amount), 2)
        rupees = int(amount)
        paise = int((amount - rupees) * 100)

        result = []
        
        if rupees == 0:
            result.append("Zero")
        else:
            crore = rupees // 10000000
            rupees %= 10000000
            lakh = rupees // 100000
            rupees %= 100000
            thousand = rupees // 1000
            rupees %= 1000

            if crore:
                result.append(f"{convert_below_thousand(crore)} Crore")
            if lakh:
                result.append(f"{convert_below_thousand(lakh)} Lakh")
            if thousand:
                result.append(f"{convert_below_thousand(thousand)} Thousand")
            if rupees:
                result.append(convert_below_thousand(rupees))

        rupees_text = " ".join(result).strip() or "Zero"
        
        paise_text = convert_below_thousand(paise) if paise else ""
        
        final_text = rupees_text
        if paise_text:
            final_text += f" and {paise_text} Paise"
        
        return f"{final_text} Only"

    @staticmethod
    def replace_text(doc, old_text, new_text):
        """Replace text in a docx document, handling text split across runs, headers, and footers."""
        def replace_in_paragraph(paragraph, old_text, new_text):
            # Collect full text from all runs
            full_text = ''
            runs = paragraph.runs
            for i, run in enumerate(runs):
                full_text += run.text

            # Check if placeholder exists in the full text
            if old_text not in full_text:
                return

            # Replace the text while preserving formatting
            remaining_text = full_text.replace(old_text, new_text, 1)  # Replace one instance at a time
            current_pos = 0

            for run in runs:
                run_text = run.text
                run_length = len(run_text)

                # If the run contains part of the old_text, replace it
                if current_pos < len(remaining_text):
                    run.text = remaining_text[current_pos:current_pos + run_length]
                    current_pos += run_length
                else:
                    run.text = ''

            # If there's remaining text, append it to the last run
            if current_pos < len(remaining_text):
                runs[-1].text += remaining_text[current_pos:]

            print(f"[DEBUG] Replaced '{old_text}' with '{new_text}' in paragraph: {full_text}")

        # Process paragraphs in the main body
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph, old_text, new_text)

        # Process paragraphs in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph, old_text, new_text)

        # Process headers and footers
        for section in doc.sections:
            for header in section.header.paragraphs:
                replace_in_paragraph(header, old_text, new_text)
            for footer in section.footer.paragraphs:
                replace_in_paragraph(footer, old_text, new_text)

    def save_invoice(self, doc):
        """Save the invoice as PDF only, using a temporary docx file."""
        save_file = filedialog.asksaveasfilename(
            defaultextension='.pdf',
            filetypes=[('PDF files', '*.pdf')]
        )
        if not save_file:
            return None

        try:
            # Create a temporary docx file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
                temp_docx_path = temp_docx.name
                doc.save(temp_docx_path)
            
            # Convert to PDF
            convert(temp_docx_path, save_file)
            
            # Clean up temporary docx file
            if os.path.exists(temp_docx_path):
                os.remove(temp_docx_path)

            # Open the PDF
            subprocess.Popen(['open', save_file])
            print(f'[+] Invoice saved to {save_file} \n {os.getcwd()}')
            messagebox.showinfo('Success', f'Invoice saved to {save_file}')
            return save_file
        except Exception as e:
            messagebox.showerror('Error', f'Failed to save or convert invoice: {e}')
            print(f'Error saving invoice: {e}')
            return None

    def create_invoice(self):
        """Create and save the invoice."""
        try:
            doc = docx.Document('TPR_template.docx')

            replacements = {
                "[Date]": self.date_entry.get(),
                "[Reference Number]": self.reference_number_entry.get(),
                "[Name]": self.name_entry.get(),
                "[Amount]": self.amount_entry.get(),
                "[Payment Mode]": self.payment_mode.get(),
                "[Payment Words]": self.amount_words_entry.get()
            }

            # Apply replacements multiple times to handle multiple instances
            for key, value in replacements.items():
                # Repeat replacement to ensure all instances are caught
                for _ in range(10):  # Arbitrary limit to avoid infinite loops
                    self.replace_text(doc, key, str(value))

            saved_file = self.save_invoice(doc)
            if saved_file:
                print(f'[+] Invoice created and saved as {saved_file}')
            else:
                print('[-] Invoice creation cancelled or failed')

        except FileNotFoundError:
            messagebox.showerror('Error', 'Template file TPR_template.docx not found.')
        except Exception as e:
            messagebox.showerror('Error', f'An error occurred: {str(e)}')

if __name__ == '__main__':
    InvoiceAutomation()